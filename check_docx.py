import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

try:
    doc = docx.Document(r'C:\Users\CNBT\Downloads\75+ Đề thi tiếng anh vào 10 HCM 2026.docx')
    c1 = sum(1 for p in doc.paragraphs if 'họ và tên' in p.text.lower())
    c2 = sum(1 for p in doc.paragraphs if 'choose the word' in p.text.lower())
    print("In Word document:")
    print("Tests by 'Họ và tên':", c1)
    print("Tests by 'Choose the word':", c2)
except Exception as e:
    print("Error:", e)
